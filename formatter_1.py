import os
import re
from docx import Document


# loop through each file in the directory

def format_transscript(directory: str, firstname: str, lastname: str):
    for filename in os.listdir(directory):
        if filename.endswith('.docx'):
            # open the document and read the content
            doc = Document(os.path.join(directory, filename))
            content = ' '
            for para in doc.paragraphs:
                content += para.text

            # replace the timestamp with a white line
            content = re.sub(r'\d+:\d+:\d+\.\d+ --> \d+:\d+:\d+\.\d+', '\n', content)

            # replace name with "researcher:"
            content = re.sub(fr"\b{firstname}\s{lastname}\b", "Researcher:", content) 

            # save the edited file as a new file
            edited_doc = Document()
            edited_doc.add_paragraph(content)
            edited_doc.save(os.path.join(directory, 'edited_' + filename))
            return 'pass'
        
        
# if __name__ == '__main__':
#     format_transscript(directory=directory)



