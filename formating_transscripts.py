import os
import re
from docx import Document

directory = "path" #! enter path of directory


# loop through each file in the directory

def format_transscript(directory):
    for filename in os.listdir(directory):
        if filename.endswith('.docx'):
            # open the document and read the content
            doc = Document(os.path.join(directory, filename))
            content = ' '
            for para in doc.paragraphs:
                content += para.text

            # replace the timestamp with a space
            content = re.sub(r'\d+:\d+:\d+\.\d+ --> \d+:\d+:\d+\.\d+', '\n', content)

            # replace my name with "researcher"
            content = re.sub(r'\bMadelief\sRiphagen\b', 'Researcher:', content)

            # save the edited file
            edited_doc = Document()
            edited_doc.add_paragraph(content)
            edited_doc.save(os.path.join(directory, 'edited_' + filename))

if __name__ == '__main__':
    format_transscript(directory=directory)



